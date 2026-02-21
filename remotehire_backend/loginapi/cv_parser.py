"""
CV Parsing and Metadata Extraction with Multi-Provider AI Support and Caching
Supports: Google Gemini, Groq, OpenAI (with fallback)
"""
import google.genai as genai
from django.conf import settings
import json
from pypdf import PdfReader
from pathlib import Path
import time
import hashlib
import os

# Try to import python-docx for DOCX support
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[CV_PARSER] python-docx not installed. DOCX support disabled.")


def extract_text_from_pdf(pdf_path):
    """Extract text from PDF file using pypdf"""
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                # Clean up any encoding issues
                page_text = page_text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')
                text += page_text
        
        if not text:
            print(f"[PDF_EXTRACT] Warning: No text extracted from PDF")
        else:
            print(f"[PDF_EXTRACT] Extracted {len(text)} characters from PDF")
        return text
    except Exception as e:
        print(f"[PDF_EXTRACT] Error extracting PDF text: {str(e)}")
        return ""


def extract_text_from_docx(docx_path):
    """Extract text from DOCX file"""
    if not DOCX_AVAILABLE:
        print("[DOCX_EXTRACT] python-docx not available")
        return ""
    
    try:
        doc = DocxDocument(docx_path)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + "\n"
        
        if not text:
            print(f"[DOCX_EXTRACT] Warning: No text extracted from DOCX")
        else:
            print(f"[DOCX_EXTRACT] Extracted {len(text)} characters from DOCX")
        return text
    except Exception as e:
        print(f"[DOCX_EXTRACT] Error extracting DOCX text: {str(e)}")
        return ""


def calculate_file_hash(file_path):
    """Calculate SHA256 hash of file for cache key"""
    try:
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        file_hash = sha256_hash.hexdigest()
        print(f"[FILE_HASH] Calculated hash: {file_hash[:16]}...")
        return file_hash
    except Exception as e:
        print(f"[FILE_HASH] Error calculating hash: {str(e)}")
        return None


def get_gemini_client():
    """Get Gemini API client using the new google.genai package"""
    api_key = getattr(settings, 'GEMINI_API_KEY', None)
    if not api_key or api_key == 'your_gemini_api_key_here':
        print(f"[GEMINI_CONFIG] API Key not configured")
        return None
    
    try:
        # The new google.genai package doesn't have configure()
        # Instead, pass api_key directly to the client
        client = genai.Client(api_key=api_key)
        print(f"[GEMINI_CONFIG] Gemini client created successfully")
        return client
    except Exception as e:
        print(f"[GEMINI_CONFIG] Error creating client: {str(e)}")
        return None


def call_ollama_api(prompt, cv_text):
    """Call local Ollama API (NO RATE LIMITS!)"""
    try:
        ollama_url = getattr(settings, 'OLLAMA_API_URL', 'http://localhost:11434')
        model = getattr(settings, 'OLLAMA_MODEL', 'gemma2')
        timeout_sec = getattr(settings, 'OLLAMA_TIMEOUT', 120)
        if isinstance(timeout_sec, str):
            try:
                timeout_sec = int(timeout_sec)
            except Exception:
                timeout_sec = 120
        num_predict = getattr(settings, 'OLLAMA_NUM_PREDICT', 500)
        if isinstance(num_predict, str):
            try:
                num_predict = int(num_predict)
            except Exception:
                num_predict = 500
        
        print(f"[OLLAMA_API] Calling local Ollama at {ollama_url} with model {model}...")
        
        import requests
        # Increase timeout to 120 seconds for large CVs and slower systems
        # Ollama can take time on first run or with large files
        response = requests.post(
            f"{ollama_url}/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.1,
                    "top_p": 0.9,
                    "num_predict": num_predict
                }
            },
            timeout=timeout_sec
        )
        
        if response.status_code == 200:
            result = response.json()
            response_text = result.get('response', '')
            if response_text:
                print(f"[OLLAMA_API] ✓ Success with local Ollama")
                return response_text
            else:
                print(f"[OLLAMA_API] Empty response from Ollama")
                return None
        else:
            print(f"[OLLAMA_API] Error {response.status_code}: {response.text[:200]}")
            return None
    except Exception as e:
        error_msg = str(e)
        if "timeout" in error_msg.lower():
            print(f"[OLLAMA_API] Request timed out. Ollama may be slow or processing large file.")
            print(f"[OLLAMA_API] Troubleshooting: 1) Ensure Ollama is running: ollama serve")
            print(f"[OLLAMA_API] Troubleshooting: 2) Check system resources")
        elif "Connection" in error_msg or "refused" in error_msg:
            print(f"[OLLAMA_API] Ollama not running or not accessible: {error_msg[:100]}")
        else:
            print(f"[OLLAMA_API] Exception: {error_msg[:100]}")
        return None


def call_gemini_api(prompt, cv_text):
    """Call Gemini API with retry logic"""
    try:
        client = get_gemini_client()
        if not client:
            return None
        
        # Retry logic for rate limits
        max_retries = 2
        retry_delay = 1
        
        for attempt in range(max_retries):
            try:
                print(f"[GEMINI_API] Attempt {attempt + 1}/{max_retries}")
                # Use the new google.genai API
                response = client.models.generate_content(
                    model="gemini-2.0-flash",
                    contents=prompt
                )
                result_text = response.text if hasattr(response, 'text') else str(response)
                if result_text:
                    print(f"[GEMINI_API] ✓ Success with Gemini API")
                    return result_text
                else:
                    print(f"[GEMINI_API] Empty response from Gemini")
                    return None
            except Exception as e:
                error_str = str(e)
                print(f"[GEMINI_API] Error: {error_str[:200]}")
                if "429" in error_str or "quota" in error_str.lower():
                    if attempt < max_retries - 1:
                        time.sleep(retry_delay)
                        retry_delay *= 2
                    else:
                        print(f"[GEMINI_API] Rate limited after retries")
                        return None
                else:
                    return None
        return None
    except Exception as e:
        print(f"[GEMINI_API] Exception: {str(e)}")
        return None


def call_groq_api(prompt, cv_text):
    """Call Groq API (free tier: 30 RPM)"""
    try:
        api_key = getattr(settings, 'GROQ_API_KEY', None)
        if not api_key or api_key == 'your_groq_api_key_here':
            print("[GROQ_API] API Key not configured")
            return None
        
        import requests
        print("[GROQ_API] Calling Groq API...")
        
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            },
            json={
                "model": "llama-3.3-70b-versatile",  # Fast and accurate
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.1
            },
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            return result['choices'][0]['message']['content']
        else:
            print(f"[GROQ_API] Error {response.status_code}: {response.text[:200]}")
            return None
    except Exception as e:
        print(f"[GROQ_API] Exception: {str(e)}")
        return None


def call_openai_api(prompt, cv_text):
    """Call OpenAI API (free trial: $5 credit)"""
    try:
        api_key = getattr(settings, 'OPENAI_API_KEY', None)
        if not api_key or api_key == 'your_openai_api_key_here':
            print("[OPENAI_API] API Key not configured")
            return None
        
        import requests
        print("[OPENAI_API] Calling OpenAI API...")
        
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            },
            json={
                "model": "gpt-3.5-turbo",  # Cheapest option
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.1
            },
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            return result['choices'][0]['message']['content']
        else:
            print(f"[OPENAI_API] Error {response.status_code}: {response.text[:200]}")
            return None
    except Exception as e:
        print(f"[OPENAI_API] Exception: {str(e)}")
        return None


def call_multi_provider_ai(prompt, cv_text):
    """Try multiple AI providers in sequence until one succeeds"""
    providers = [
        ("Ollama (Local)", call_ollama_api),  # Try local first - NO RATE LIMITS!
        ("Groq", call_groq_api),              # Groq first (fast & free)
        ("Gemini", call_gemini_api),          # Gemini as fallback
        ("OpenAI", call_openai_api)
    ]
    
    print(f"[MULTI_AI] Trying {len(providers)} providers...")
    
    for provider_name, provider_func in providers:
        try:
            print(f"[MULTI_AI] Attempting {provider_name}...")
            response_text = provider_func(prompt, cv_text)
            if response_text:
                print(f"[MULTI_AI] ✓ Success with {provider_name}")
                return response_text, provider_name
        except Exception as e:
            print(f"[MULTI_AI] {provider_name} failed: {str(e)}")
            continue
    
    print("[MULTI_AI] All providers failed")
    return None, None


def configure_gemini_old():
    """Configure Gemini API using the new google.genai package"""
    api_key = getattr(settings, 'GEMINI_API_KEY', None)
    print(f"[GEMINI_CONFIG] API Key configured: {bool(api_key)}")
    if not api_key or api_key == 'your_gemini_api_key_here':
        raise ValueError("GEMINI_API_KEY not configured in .env")
    
    # Configure with the new google.genai client
    genai.configure(api_key=api_key)
    
    # List available models
    try:
        models = genai.list_models()
        print(f"[GEMINI_CONFIG] Available models:")
        for model in models:
            if 'generateContent' in model.supported_generation_methods:
                print(f"  - {model.name}")
    except Exception as e:
        print(f"[GEMINI_CONFIG] Could not list models: {str(e)}")
    
    print(f"[GEMINI_CONFIG] Gemini configured successfully")


def extract_cv_metadata(cv_file_path, cached_metadata=None, file_hash=None):
    """
    Extract metadata from CV using Multi-Provider AI with Caching
    Returns structured metadata including skills, experience, education, etc.
    
    Args:
        cv_file_path: Path to the CV file
        cached_metadata: Previously cached metadata (if exists)
        file_hash: Hash of the current file
    """
    try:
        # Calculate file hash for cache comparison
        if not file_hash:
            file_hash = calculate_file_hash(cv_file_path)
        
        # Check if we have valid cached metadata
        if cached_metadata and isinstance(cached_metadata, dict):
            cached_hash = cached_metadata.get('_file_hash')
            if cached_hash == file_hash and not cached_metadata.get('error'):
                print(f"[CV_METADATA] ✓ Using cached metadata (hash match)")
                return cached_metadata
            else:
                print(f"[CV_METADATA] Cache invalid or file changed, re-analyzing...")
        
        # Extract text based on file type
        file_ext = str(cv_file_path).lower()
        print(f"[CV_METADATA] Processing file: {cv_file_path}")
        
        if file_ext.endswith('.pdf'):
            cv_text = extract_text_from_pdf(cv_file_path)
        elif file_ext.endswith('.docx'):
            cv_text = extract_text_from_docx(cv_file_path)
        elif file_ext.endswith('.txt'):
            try:
                with open(cv_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    cv_text = f.read()
                print(f"[TXT_EXTRACT] Extracted {len(cv_text)} characters from TXT")
            except Exception as e:
                print(f"[TXT_EXTRACT] Error: {str(e)}")
                cv_text = ""
        else:
            print(f"[CV_METADATA] Unsupported file type: {file_ext}")
            return {"error": f"Unsupported file type. Supported: PDF, DOCX, TXT"}
        
        if not cv_text.strip():
            return {"error": "Could not extract text from CV. The file may be empty or corrupted."}
        
        print(f"[CV_METADATA] Extracted CV text length: {len(cv_text)} characters")
        
        # Trim excessively long CV text to avoid timeouts in local models
        max_chars = getattr(settings, 'CV_TEXT_MAX_CHARS', 12000)
        if isinstance(max_chars, str):
            try:
                max_chars = int(max_chars)
            except Exception:
                max_chars = 12000
        if len(cv_text) > max_chars:
            print(f"[CV_METADATA] Trimming CV text from {len(cv_text)} to {max_chars} characters to improve performance")
            cv_text = cv_text[:max_chars]

        # Build prompt for structured data extraction
        prompt = f"""
        Analyze the following CV/Resume and extract the following information in JSON format:
        
        {{
            "full_name": "Extracted full name or null",
            "email": "Extracted email or null",
            "phone": "Extracted phone number or null",
            "current_title": "Current job title or null",
            "years_of_experience": numeric value (total years),
            "skills": ["list", "of", "key", "technical", "skills"],
            "programming_languages": ["list", "of", "programming", "languages"],
            "frameworks": ["list", "of", "frameworks", "and", "libraries"],
            "education": [
                {{
                    "degree": "Bachelor/Master/PhD etc",
                    "field": "Field of study",
                    "institution": "University/College name"
                }}
            ],
            "work_experience": [
                {{
                    "title": "Job title",
                    "company": "Company name",
                    "years": number (duration in years),
                    "key_achievements": ["achievement1", "achievement2"]
                }}
            ],
            "certifications": ["list", "of", "certifications"],
            "languages": ["list", "of", "languages"],
            "specializations": ["list", "of", "specialized", "areas"],
            "soft_skills": ["list", "of", "soft", "skills"]
        }}
        
        CV Content:
        {cv_text}
        
        Return ONLY valid JSON, no markdown or additional text.
        """
        
        print(f"[CV_METADATA] Calling multi-provider AI...")
        
        # Try multiple AI providers with fallback
        response_text, provider_used = call_multi_provider_ai(prompt, cv_text)
        
        if not response_text:
            return {
                "error": "All AI services are currently unavailable. Please try again later.",
                "years_of_experience": 0,
                "skills": [],
                "education": [],
                "work_experience": []
            }
        
        print(f"[CV_PARSER] Response from {provider_used}, length: {len(response_text)}")
        print(f"[CV_PARSER] Response (first 500 chars): {response_text[:500]}")
        
        # Try to parse JSON response
        try:
            # Handle markdown code blocks if present
            if response_text.startswith("```"):
                response_text = response_text.split("```")[1]
                if response_text.startswith("json"):
                    response_text = response_text[4:]
            
            # Clean up the response text - remove any BOM or invalid characters
            response_text = response_text.encode('utf-8', errors='ignore').decode('utf-8')
            response_text = response_text.strip()
            
            print(f"[CV_PARSER] After cleanup, response: {response_text[:500]}")
            
            # Try to parse JSON
            metadata = json.loads(response_text)
            print(f"[CV_PARSER] Successfully parsed metadata with keys: {list(metadata.keys())}")
            
            # Add file hash and provider info for caching
            metadata['_file_hash'] = file_hash
            metadata['_ai_provider'] = provider_used
            
            return metadata
        except json.JSONDecodeError as e:
            print(f"[CV_PARSER] JSON parsing error: {str(e)}")
            print(f"[CV_PARSER] Response that failed to parse: {response_text[:500]}")
            # Return a basic default structure instead of raw response
            return {
                "error": "AI analysis could not parse the CV format. Please ensure your CV is properly formatted with clear sections (Experience, Education, Skills, etc.).",
                "years_of_experience": 0,
                "skills": [],
                "education": [],
                "work_experience": []
            }
            
    except ValueError as e:
        return {"error": str(e)}
    except Exception as e:
        error_str = str(e)
        print(f"[CV_METADATA] CV extraction error: {error_str}")
        
        # Check if it's a rate limit error
        if "429" in error_str or "quota" in error_str.lower():
            return {
                "error": "AI service is currently at capacity. Please try uploading your CV again in a few moments. If this persists, please contact support.",
                "years_of_experience": 0,
                "skills": [],
                "education": [],
                "work_experience": []
            }
        
        return {
            "error": f"Failed to extract CV metadata: {error_str[:100]}",
            "years_of_experience": 0,
            "skills": [],
            "education": [],
            "work_experience": []
        }


def calculate_similarity_score(cv_metadata, job_requirements):
    """
    Calculate similarity score between CV metadata and job requirements
    Returns score 0-100
    """
    if not cv_metadata or not job_requirements:
        return 0.0
    
    if "error" in cv_metadata:
        return 0.0
    
    score = 0.0
    weights = {
        'skills': 0.35,
        'experience': 0.25,
        'education': 0.15,
        'languages': 0.10,
        'certifications': 0.15
    }
    
    # Skills matching
    cv_skills = set(str(s).lower() for s in cv_metadata.get('skills', []))
    required_skills = set(str(s).lower() for s in job_requirements.get('required_skills', []))
    
    if required_skills:
        skill_match = len(cv_skills & required_skills) / len(required_skills)
        score += skill_match * weights['skills'] * 100
    
    # Experience matching
    cv_experience = cv_metadata.get('years_of_experience', 0)
    required_experience = job_requirements.get('required_experience_years', 0)
    
    if required_experience > 0:
        exp_match = min(cv_experience / required_experience, 1.0)
        score += exp_match * weights['experience'] * 100
    else:
        score += weights['experience'] * 100
    
    # Education matching
    cv_education_levels = {e.get('degree', '').lower() for e in cv_metadata.get('education', [])}
    required_education = job_requirements.get('required_education', '').lower()
    
    education_match = 0.0
    if required_education:
        if required_education in cv_education_levels:
            education_match = 1.0
        elif 'bachelor' in required_education and any('bachelor' in e for e in cv_education_levels):
            education_match = 0.8
        elif 'master' in required_education and 'master' in cv_education_levels:
            education_match = 1.0
    else:
        education_match = 0.6  # Default if no education requirement specified
    
    score += education_match * weights['education'] * 100
    
    # Languages
    cv_languages = set(str(l).lower() for l in cv_metadata.get('languages', []))
    required_languages = set(str(l).lower() for l in job_requirements.get('required_languages', []))
    
    if required_languages:
        lang_match = len(cv_languages & required_languages) / len(required_languages)
        score += lang_match * weights['languages'] * 100
    
    # Certifications
    cv_certifications = set(str(c).lower() for c in cv_metadata.get('certifications', []))
    required_certifications = set(str(c).lower() for c in job_requirements.get('required_certifications', []))
    
    if required_certifications:
        cert_match = len(cv_certifications & required_certifications) / len(required_certifications)
        score += cert_match * weights['certifications'] * 100
    
    return min(score, 100.0)
